"""Render the paper as a .docx.

Reads the built HTML and walks it, so the Word file and the web version stay the
same document rather than two drafts that drift. Figures are SVG, which Word
renders inconsistently, so each figure is rasterised to PNG where a rasteriser is
available and otherwise replaced by a labelled placeholder pointing at the table
that carries the same numbers -- every figure in this paper has one.
"""
import os as _os
_HERE = _os.path.dirname(_os.path.abspath(__file__))
_REPO = _os.path.dirname(_os.path.dirname(_HERE))
import os
import re
from html import unescape
from html.parser import HTMLParser

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor, Inches

S = _HERE          # figures + template live beside these scripts
SRC = _os.path.join(_HERE, 'zero-trust-sdn-paper.html')   # intermediate, not a deliverable
OUT = _os.path.join(_HERE, '..', '..', 'docs', 'paper', 'zero-trust-sdn-paper.docx')

ACCENT = RGBColor(0x1F, 0x5F, 0xA8)
INK = RGBColor(0x12, 0x15, 0x1A)
INK2 = RGBColor(0x4A, 0x51, 0x5C)
MUTED = RGBColor(0x76, 0x7E, 0x8B)
OK = RGBColor(0x1D, 0x7A, 0x4C)
WARN = RGBColor(0xA3, 0x5A, 0x00)


# --------------------------------------------------------------------------- #
# A tiny inline-run model: (text, bold, italic, mono, color)                    #
# --------------------------------------------------------------------------- #
class Inline(HTMLParser):
    """Flatten an HTML fragment into styled runs."""

    def __init__(self):
        super().__init__(convert_charrefs=True)
        self.runs = []
        self.b = self.i = self.m = 0
        self.color = None

    def handle_starttag(self, tag, attrs):
        a = dict(attrs)
        if tag in ('b', 'strong'):
            self.b += 1
        elif tag in ('i', 'em'):
            self.i += 1
        elif tag == 'code':
            self.m += 1
        elif tag == 'span':
            cls = a.get('class', '')
            if 'm' == cls or cls.startswith('m '):
                self.m += 1
            elif 'pill' in cls:
                self.m += 1
        elif tag == 'sub':
            pass
        elif tag == 'br':
            self.runs.append(('\n', 0, 0, 0, None))

    def handle_endtag(self, tag):
        if tag in ('b', 'strong'):
            self.b = max(0, self.b - 1)
        elif tag in ('i', 'em'):
            self.i = max(0, self.i - 1)
        elif tag == 'code':
            self.m = max(0, self.m - 1)
        elif tag == 'span':
            self.m = max(0, self.m - 1)

    def handle_data(self, d):
        if d:
            self.runs.append((d, self.b, self.i, self.m, self.color))


def inline_runs(frag):
    p = Inline()
    p.feed(frag)
    out = []
    for t, b, i, m, c in p.runs:
        t = re.sub(r'\s+', ' ', t)
        if t.strip() or (out and not out[-1][0].endswith(' ')):
            out.append((t, bool(b), bool(i), bool(m), c))
    return out


def emit(par, frag, size=10.5, color=None, base_bold=False):
    for t, b, i, m, _ in inline_runs(frag):
        if not t:
            continue
        r = par.add_run(t)
        r.bold = b or base_bold
        r.italic = i
        r.font.size = Pt(size - (0.6 if m else 0))
        r.font.name = 'Consolas' if m else 'Palatino Linotype'
        if m:
            r.font.color.rgb = INK2
        elif color:
            r.font.color.rgb = color
    return par


def strip_tags(frag):
    return re.sub(r'\s+', ' ', re.sub(r'<[^>]+>', '', frag)).strip()


def shade(cell, hexfill):
    el = OxmlElement('w:shd')
    el.set(qn('w:val'), 'clear')
    el.set(qn('w:fill'), hexfill)
    cell._tc.get_or_add_tcPr().append(el)


def set_repeat_header(row):
    tr = row._tr
    pr = tr.get_or_add_trPr()
    el = OxmlElement('w:tblHeader')
    el.set(qn('w:val'), 'true')
    pr.append(el)


# --------------------------------------------------------------------------- #
def build():
    html = open(SRC).read()
    body = html.split('<div class="wrap">', 1)[1]

    doc = Document()

    st = doc.styles['Normal']
    st.font.name = 'Palatino Linotype'
    st.font.size = Pt(10.5)
    st.paragraph_format.space_after = Pt(7)
    st.paragraph_format.line_spacing = 1.15

    for s in doc.sections:
        s.left_margin = s.right_margin = Inches(1.0)
        s.top_margin = s.bottom_margin = Inches(0.9)

    # ---------------- masthead ----------------
    title = strip_tags(re.search(r'<h1 class="title">(.*?)</h1>', body, re.S).group(1))
    sub = strip_tags(re.search(r'<p class="subtitle">(.*?)</p>', body, re.S).group(1))
    eyebrow = strip_tags(re.search(r'<p class="eyebrow">(.*?)</p>', body, re.S).group(1))

    p = doc.add_paragraph()
    r = p.add_run(eyebrow.upper())
    r.font.size = Pt(7.5)
    r.font.name = 'Segoe UI'
    r.font.color.rgb = MUTED
    r.bold = True

    p = doc.add_paragraph()
    r = p.add_run(title)
    r.font.size = Pt(23)
    r.bold = True
    r.font.color.rgb = INK
    p.paragraph_format.space_after = Pt(4)

    p = doc.add_paragraph()
    r = p.add_run(sub)
    r.font.size = Pt(12)
    r.italic = True
    r.font.color.rgb = INK2
    p.paragraph_format.space_after = Pt(10)

    for line in re.findall(r'<span><b>(.*?)</b>(.*?)</span>', body, re.S):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(1)
        r = p.add_run(strip_tags(line[0]) + ' ')
        r.bold = True
        r.font.size = Pt(9)
        r = p.add_run(strip_tags(line[1]))
        r.font.size = Pt(9)
        r.font.color.rgb = INK2

    doc.add_paragraph()

    # ---------------- abstract ----------------
    abs_html = re.search(r'<section class="abstract">(.*?)</section>', body, re.S).group(1)
    p = doc.add_paragraph()
    r = p.add_run('ABSTRACT')
    r.bold = True
    r.font.size = Pt(8)
    r.font.name = 'Segoe UI'
    r.font.color.rgb = ACCENT
    for para in re.findall(r'<p>(.*?)</p>', abs_html, re.S):
        q = doc.add_paragraph()
        q.paragraph_format.left_indent = Inches(0.25)
        q.paragraph_format.space_after = Pt(6)
        emit(q, para, size=10)

    doc.add_page_break()

    # ---------------- body walk ----------------
    body = body.split('</section>', 1)[1]
    fignum = {'FIG1': 1, 'FIG2': 2, 'FIG3': 3, 'FIG4': 4, 'FIG5': 5}

    token = re.compile(
        r'<h2[^>]*>(?P<h2>.*?)</h2>'
        r'|<h3[^>]*>(?P<h3>.*?)</h3>'
        r'|<h4[^>]*>(?P<h4>.*?)</h4>'
        r'|<div class="keyfinding">(?P<kf>.*?)</div>\s*</div>'
        r'|<div class="note">(?P<note>.*?)</div>'
        r'|<div class="eq">(?P<eq>.*?)</div>'
        r'|<pre>(?P<pre>.*?)</pre>'
        r'|<figure>(?P<fig>.*?)</figure>'
        r'|<div class="tbl-wrap">\s*(?P<tbl><table>.*?</table>)\s*</div>'
        r'|<ol class="contribs">(?P<contribs>.*?)</ol>'
        r'|<ol class="refs">(?P<refs>.*?)</ol>'
        r'|<ol>(?P<ol>.*?)</ol>'
        r'|<ul>(?P<ul>.*?)</ul>'
        r'|<p[^>]*>(?P<p>.*?)</p>'
        r'|<footer class="colophon">(?P<foot>.*?)</footer>',
        re.S)

    for m in token.finditer(body):
        g = m.lastgroup
        v = m.group(g)

        if g in ('h2', 'h3', 'h4'):
            num = re.search(r'<span class="num">(.*?)</span>', v)
            txt = strip_tags(re.sub(r'<span class="num">.*?</span>', '', v))
            label = (f"{strip_tags(num.group(1))}  " if num else '') + txt
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(16 if g == 'h2' else 11)
            p.paragraph_format.space_after = Pt(4)
            p.paragraph_format.keep_with_next = True
            r = p.add_run(label)
            r.bold = True
            r.font.color.rgb = INK if g != 'h4' else INK2
            r.font.size = Pt({'h2': 15, 'h3': 12, 'h4': 9.5}[g])
            if g == 'h4':
                r.font.name = 'Segoe UI'
            if g == 'h2':
                pr = p._p.get_or_add_pPr()
                bd = OxmlElement('w:pBdr')
                bot = OxmlElement('w:bottom')
                bot.set(qn('w:val'), 'single')
                bot.set(qn('w:sz'), '6')
                bot.set(qn('w:color'), 'C3CAD4')
                bd.append(bot)
                pr.append(bd)

        elif g == 'kf':
            lab = re.search(r'<p class="kf-label">(.*?)</p>', v)
            if lab:
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Inches(0.3)
                p.paragraph_format.space_after = Pt(2)
                r = p.add_run(strip_tags(lab.group(1)).upper())
                r.bold = True
                r.font.size = Pt(8)
                r.font.name = 'Segoe UI'
                r.font.color.rgb = ACCENT
            for para in re.findall(r'<p>(.*?)</p>', v, re.S):
                q = doc.add_paragraph()
                q.paragraph_format.left_indent = Inches(0.3)
                q.paragraph_format.space_after = Pt(5)
                emit(q, para, size=10)

        elif g == 'note':
            for para in re.findall(r'<p>(.*?)</p>', v, re.S):
                q = doc.add_paragraph()
                q.paragraph_format.left_indent = Inches(0.3)
                emit(q, para, size=9.5, color=INK2)

        elif g in ('eq', 'pre'):
            # These blocks are preformatted, so they bypass the inline parser and
            # must unescape entities themselves -- `&lt;` in a comparison operator
            # otherwise reaches the page literally.
            for line in v.replace('<span class="cmt">', '').replace('</span>', '').split('\n'):
                line = strip_tags(line) if '<' in line else line
                line = unescape(line)
                q = doc.add_paragraph()
                q.paragraph_format.space_after = Pt(0)
                q.paragraph_format.left_indent = Inches(0.3)
                r = q.add_run(line.rstrip() or ' ')
                r.font.name = 'Consolas'
                r.font.size = Pt(8.5)
                r.font.color.rgb = INK2
            doc.add_paragraph().paragraph_format.space_after = Pt(2)

        elif g == 'fig':
            cap = re.search(r'<figcaption>(.*?)</figcaption>', v, re.S)
            n = re.search(r'<b>Figure (\d+)\.</b>', cap.group(1)) if cap else None
            png = f"{S}/fig{n.group(1)}.png" if n else None
            q = doc.add_paragraph()
            q.alignment = WD_ALIGN_PARAGRAPH.CENTER
            q.paragraph_format.space_before = Pt(10)
            q.paragraph_format.space_after = Pt(4)
            q.paragraph_format.keep_with_next = True
            if png and os.path.exists(png):
                q.add_run().add_picture(png, width=Inches(6.5))
            else:
                r = q.add_run(f"[ Figure {n.group(1) if n else '?'} unavailable ]")
                r.italic = True
                r.font.size = Pt(9)
                r.font.color.rgb = MUTED
            if cap:
                q = doc.add_paragraph()
                q.paragraph_format.space_after = Pt(12)
                emit(q, cap.group(1), size=9, color=INK2)

        elif g == 'tbl':
            add_table(doc, v)

        elif g in ('contribs', 'refs', 'ol', 'ul'):
            items = re.findall(r'<li[^>]*>(.*?)</li>', v, re.S)
            for idx, it in enumerate(items, 1):
                q = doc.add_paragraph()
                q.paragraph_format.left_indent = Inches(0.42)
                q.paragraph_format.first_line_indent = Inches(-0.42)
                q.paragraph_format.space_after = Pt(3)
                marker = {'contribs': f'C{idx}', 'refs': f'[{idx}]',
                          'ol': f'{idx}.', 'ul': '•'}[g]
                r = q.add_run(marker + '\t')
                r.font.name = 'Consolas'
                r.font.size = Pt(9)
                r.font.color.rgb = ACCENT if g in ('contribs', 'refs') else MUTED
                emit(q, it, size=10 if g != 'refs' else 9)

        elif g == 'p':
            if 'class="eyebrow"' in m.group(0) or 'class="subtitle"' in m.group(0):
                continue
            txt = strip_tags(v)
            if not txt:
                continue
            q = doc.add_paragraph()
            q.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            emit(q, v)

        elif g == 'foot':
            doc.add_paragraph()
            p = doc.add_paragraph()
            r = p.add_run('PROVENANCE')
            r.bold = True
            r.font.size = Pt(8)
            r.font.name = 'Segoe UI'
            r.font.color.rgb = MUTED
            for para in re.findall(r'<p>(.*?)</p>', v, re.S):
                q = doc.add_paragraph()
                emit(q, para, size=9, color=INK2)

    doc.save(OUT)
    return OUT


def add_table(doc, tbl_html):
    cap = re.search(r'<caption>(.*?)</caption>', tbl_html, re.S)
    if cap:
        q = doc.add_paragraph()
        q.paragraph_format.space_before = Pt(8)
        q.paragraph_format.space_after = Pt(3)
        q.paragraph_format.keep_with_next = True
        emit(q, cap.group(1), size=9, color=INK2)

    head = re.search(r'<thead>(.*?)</thead>', tbl_html, re.S)
    bodym = re.search(r'<tbody>(.*?)</tbody>', tbl_html, re.S)
    hrows = re.findall(r'<tr[^>]*>(.*?)</tr>', head.group(1), re.S) if head else []
    brows = re.findall(r'<tr[^>]*>(.*?)</tr>', bodym.group(1), re.S) if bodym else []

    def cells(tr):
        return re.findall(r'<t[hd]([^>]*)>(.*?)</t[hd]>', tr, re.S)

    ncol = 0
    for tr in hrows + brows:
        n = 0
        for attrs, _ in cells(tr):
            cs = re.search(r'colspan="(\d+)"', attrs)
            n += int(cs.group(1)) if cs else 1
        ncol = max(ncol, n)
    if not ncol:
        return

    t = doc.add_table(rows=0, cols=ncol)
    t.style = 'Table Grid'
    t.alignment = WD_TABLE_ALIGNMENT.CENTER

    for ri, tr in enumerate(hrows + brows):
        is_head = ri < len(hrows)
        row = t.add_row()
        if is_head:
            set_repeat_header(row)
        ci = 0
        for attrs, html_cell in cells(tr):
            if ci >= ncol:
                break
            cell = row.cells[ci]
            span = re.search(r'colspan="(\d+)"', attrs)
            width = int(span.group(1)) if span else 1
            if width > 1 and ci + width <= ncol:
                cell = cell.merge(row.cells[ci + width - 1])
            cell.text = ''
            par = cell.paragraphs[0]
            par.paragraph_format.space_after = Pt(1)
            par.paragraph_format.space_before = Pt(1)
            color = None
            if 'class="win"' in attrs:
                color = OK
            elif 'class="lose"' in attrs:
                color = WARN
            emit(par, html_cell, size=8.5, color=color,
                 base_bold=is_head or 'best' in attrs)
            if is_head:
                shade(cell, 'E9EDF2')
            elif 'sysrow' in tr:
                shade(cell, 'E6EEF8')
            ci += width
    doc.add_paragraph().paragraph_format.space_after = Pt(4)


if __name__ == '__main__':
    print('wrote', build())
